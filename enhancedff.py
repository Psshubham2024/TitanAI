import streamlit as st
import requests
from docx import Document
import PyPDF2
import json
from io import BytesIO
import re

# Define the API endpoint and access token
API_URL = "https://api.psnext.info/api/chat"
PSCHATACCESSTOKEN = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJVc2VySW5mbyI6eyJpZCI6MzcxMzcsInJvbGVzIjpbImRlZmF1bHQiXSwicGF0aWQiOiJmYTRlMzZhNC1mMWU5LTRjMjktODYwMi0wZDU1NGFmMGIxYzcifSwiaWF0IjoxNzI2MDUwNDI2LCJleHAiOjE3Mjg2NDI0MjZ9.gw74MhUO6rO3wrauMHQxdm8PWK6RBJAh6v7yIFSS8zA"  # Replace with your actual token

# Function to call PSNext API for CV matching and rating
def get_cv_match(cv_text, job_description):
    payload = {
        "message": f"Evaluate this CV against the following job description, providing a rating out of 10 and feedback:\n\nJob Description:\n{job_description}\n\nCV:\n{cv_text}",
        "options": {"model": "gpt35turbo"}
    }
    
    headers = {
        "Authorization": f"Bearer {PSCHATACCESSTOKEN}",
        "Content-Type": "application/json"
    }
    
    response = requests.post(API_URL, headers=headers, data=json.dumps(payload))
    
    if response.status_code == 200:
        response_data = response.json()
        messages = response_data.get('data', {}).get('messages', [])
        for message in messages:
            if message.get('role') == 'assistant':
                result = message.get('content', 'No content returned from the API.')
                # Extract the rating and feedback from the result
                match = re.search(r'Rating: (\d+)/10', result)
                if match:
                    rating = int(match.group(1))
                    feedback = result.split('\n', 1)[1] if '\n' in result else ''
                    return rating, feedback
                else:
                    return 0, 'Unable to extract rating from the response.'
        return 0, 'No assistant message found in the API response.'
    else:
        return 0, f"Error: {response.status_code}, {response.text}"

# Function to generate case study questions
def generate_case_study_questions(job_description, years_of_experience, industry, difficulty_level):
    payload = {
        "message": f"Generate a set of case study questions based on the following job description, {years_of_experience} years of experience in the {industry} industry, with a difficulty level of {difficulty_level}:\n\nJob Description:\n{job_description}",
        "options": {"model": "gpt35turbo"}
    }
    
    headers = {
        "Authorization": f"Bearer {PSCHATACCESSTOKEN}",
        "Content-Type": "application/json"
    }
    
    response = requests.post(API_URL, headers=headers, data=json.dumps(payload))
    
    if response.status_code == 200:
        response_data = response.json()
        messages = response_data.get('data', {}).get('messages', [])
        for message in messages:
            if message.get('role') == 'assistant':
                return message.get('content', 'No content returned from the API.')
        return 'No assistant message found in the API response.'
    else:
        return f"Error: {response.status_code}, {response.text}"

# Function to match case study answers
def match_case_study_answers(question, provided_answer):
    payload = {
        "message": f"Evaluate the following case study question and answer, providing a rating out of 10 and a brief feedback:\n\nQuestion:\n{question}\n\nAnswer:\n{provided_answer}",
        "options": {"model": "gpt35turbo"}
    }
    
    headers = {
        "Authorization": f"Bearer {PSCHATACCESSTOKEN}",
        "Content-Type": "application/json"
    }
    
    response = requests.post(API_URL, headers=headers, data=json.dumps(payload))
    
    if response.status_code == 200:
        response_data = response.json()
        messages = response_data.get('data', {}).get('messages', [])
        for message in messages:
            if message.get('role') == 'assistant':
                result = message.get('content', 'No content returned from the API.')
                # Extract the rating and feedback from the result
                match = re.search(r'Rating: (\d+)/10', result)
                if match:
                    rating = int(match.group(1))
                    feedback = result.split('\n', 1)[1] if '\n' in result else ''
                    return rating, feedback
                else:
                    return 0, 'Unable to extract rating from the response.'
        return 0, 'No assistant message found in the API response.'
    else:
        return 0, f"Error: {response.status_code}, {response.text}"

# Function to extract text from a Word document
def extract_text_from_word(docx_file):
    doc = Document(docx_file)
    return '\n'.join([para.text for para in doc.paragraphs])

# Function to extract text from a PDF file
def extract_text_from_pdf(pdf_file):
    reader = PyPDF2.PdfReader(pdf_file)
    cv_text = ""
    for page_num in range(len(reader.pages)):
        page = reader.pages[page_num]
        cv_text += page.extract_text()
    return cv_text

# Function to create a Word document with case study questions
def create_case_study_doc(case_study_text):
    doc = Document()
    doc.add_heading('Generated Case Study Questions', 0)
    doc.add_paragraph(case_study_text)
    
    # Save to a BytesIO object for download
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Emoji mapping based on rating
def get_rating_emoji(rating):
    if rating >= 9:
        return "🌟 Excellent"
    elif rating >= 7:
        return "👍 Good"
    elif rating >= 5:
        return "👌 Average"
    elif rating >= 3:
        return "🤔 Okay"
    else:
        return "👎 Bad"

# Welcome page
def welcome_page():
    st.title("CareerQgen-Your AI-Powered Staffing Solution!")
    st.subheader("Optimize Your Hiring Process with AI ©TitanAI")
    st.markdown("""
        This application uses cutting-edge AI technology to evaluate candidate CVs against job descriptions,
        and generate tailored case study questions for interviews. 
        
        **Features:**
        - CV matching for job requirements
        - Generate case study questions based on job descriptions and experience of the candidate
        - Match case study answers with provided responses
        
        
        Please login to continue.
        
    """)
    if st.button("Continue"):
        st.session_state["page"] = "login"

# Login page
def login_page():
    st.title("Login")
    username = st.text_input("Username", placeholder="Enter your username")
    password = st.text_input("Password", type="password", placeholder="Enter your password")
    
    if st.button("Login"):
        if username == "Titanai" and password == "Password":
            st.session_state["logged_in"] = True
            st.session_state["page"] = "main"
        else:
            st.error("Invalid username or password")

# Main app page
def main_app():
    st.title("CareerQgen-Your AI-Powered Staffing Solution")
    st.subheader("Optimize Your Hiring Process with AI")

    # Add a logout button
    if st.button("Logout"):
        st.session_state["logged_in"] = False
        st.session_state["page"] = "welcome"

    tabs = st.tabs(["Profile Matching & Case Study Generation", "Case Study evaluation"])

    # Tab 1: CV Matching & Case Study Questions
    with tabs[0]:
        st.header("CV Matching & Case Study evaluation")
        st.write("Upload a CV (PDF or Word) and provide a job description to get a match evaluation and case study questions.")

        # Upload CV (Word or PDF)
        uploaded_cv = st.file_uploader("Upload Candidate CV (PDF or Word)", type=["pdf", "docx"])
        job_description = st.text_area("Job Description", height=200)
        years_of_experience = st.number_input("Years of Experience", min_value=0, max_value=50, step=1)
        industry = st.selectbox("Industry/Domain", ["Finance", "Supply Chain", "Technology", "Healthcare", "Education", "Other"])
        difficulty_level = st.selectbox("Difficulty Level", ["Easy", "Moderate", "Difficult"])

        result_output = st.empty()
        case_study_output = st.empty()

        # Process CV file and match
        if uploaded_cv and job_description and years_of_experience:
            if uploaded_cv.type == "application/pdf":
                cv_text = extract_text_from_pdf(uploaded_cv)
            elif uploaded_cv.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                cv_text = extract_text_from_word(uploaded_cv)

            # Ensure CV text is not empty
            if cv_text.strip():
                if st.button("Get Match Analysis"):
                    with st.spinner("Analyzing CV..."):
                        rating, result = get_cv_match(cv_text, job_description)
                        result_output.text_area("Match Analysis Result", result, height=200)
                        st.write(f"Rating: {rating}/10 {get_rating_emoji(rating)}")

                if st.button("Generate Case Study Questions"):
                    with st.spinner("Generating questions..."):
                        case_study_result = generate_case_study_questions(job_description, years_of_experience, industry, difficulty_level)
                        case_study_output.text_area("Case Study Questions", case_study_result, height=200)

                        # Allow user to download the generated case study as a .docx file
                        if case_study_result:
                            case_study_doc = create_case_study_doc(case_study_result)
                            st.download_button(
                                label="Download Case Study as DOCX",
                                data=case_study_doc,
                                file_name="case_study_questions.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
            else:
                st.error("The uploaded CV appears to be empty. Please upload a valid CV.")

    # Tab 2: Case Study Answer Matching
    with tabs[1]:
        st.header("Case Study Response evaluation")
        st.write("Match case study questions with their provided answers.")

        question = st.text_area("Case Study Question", height=100)
        provided_answer = st.text_area("Provided Answer", height=200)

        answer_result = st.empty()

        if st.button("Match Case Study Answer"):
            if question and provided_answer:
                with st.spinner("Matching answer..."):
                    rating, feedback = match_case_study_answers(question, provided_answer)
                    answer_result.text_area("Answer Matching Result", feedback, height=200)
                    st.write(f"Rating: {rating}/10 {get_rating_emoji(rating)}")
            else:
                st.error("Please provide both a question and an answer.")

# Main logic to handle page navigation
if "page" not in st.session_state:
    st.session_state["page"] = "welcome"

if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False

if st.session_state["page"] == "welcome":
    welcome_page()
elif st.session_state["page"] == "login":
    login_page()
elif st.session_state["logged_in"]:
    main_app()
else:
    st.error("You must log in to access the application.")
